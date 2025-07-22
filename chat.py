import streamlit as st
import smtplib
import random
import textwrap
import string
import json
import os
import io
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from PyPDF2 import PdfReader
import docx
from gtts import gTTS
from langdetect import detect
from deep_translator import GoogleTranslator
from llama_index.core import VectorStoreIndex, StorageContext, Document
from llama_index.llms.together import TogetherLLM
from llama_index.core.node_parser import SentenceSplitter
from langchain_community.embeddings import HuggingFaceInferenceAPIEmbeddings
from together import Together
from llama_index.core import Settings
from langchain_community.embeddings import HuggingFaceHubEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings
# st.set_page_config(page_title="Document Chatbot", layout="wide")

LANGUAGE_CODES = {
    'en': 'English', 'hi': 'Hindi', 'mr': 'Marathi', 'bn': 'Bengali',
    'gu': 'Gujarati', 'te': 'Telugu', 'ta': 'Tamil', 'ml': 'Malayalam',
    'kn': 'Kannada', 'pa': 'Punjabi', 'or': 'Odia'
}
# Custom CSS to style the page
st.markdown("""
    <style>
    .main {
        background-color: #f0f0f5;
        color: #4c4c8a;
    }
    .stButton>button {
        background-color: #4c4c8a;
        color: white;
        border-radius: 10px;
        padding: 10px;
    }
    .stTextInput>div>div>input {
        border: 2px solid #4c4c8a;
        padding: 10px;
    }
    .stSidebar {
        background-color: #333366;
        color: white;
    }
    .stSidebar button {
        background-color: #4c4c8a;
        color: white;
        border-radius: 5px;
    }
    </style>
    """, unsafe_allow_html=True)


# Function to generate OTP
def generate_otp(length=6):
    digits = string.digits
    otp = ''.join(random.choice(digits) for _ in range(length))
    return otp

# Function to send OTP email
def send_otp_email(sender_email, app_password, receiver_email, otp):
    subject = "Your OTP Code"
    body = f"Your OTP code is {otp}. Please use this code to complete your verification."
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, app_password)
        server.sendmail(sender_email, receiver_email, msg.as_string())
        server.quit()
        st.success(f"OTP sent successfully to {receiver_email}")
    except Exception as e:
        st.error(f"Failed to send OTP. Error: {str(e)}")

# Function to verify OTP
def verify_otp(user_input, otp):
    return user_input == otp

@st.cache_data
def load_login_data():
    with open(r'./data.json') as f:
        data = json.load(f)
    return data

def login():
    st.title("🔑 Login")
    st.markdown("### Please enter your credentials")
    username = st.text_input("📧 Email")
    password = st.text_input("🔒 Password", type="password")
    
    if st.button("Login"):
        data = load_login_data()
        user_found = False
        for user in data['login']:
            if username == user['email'] and password == user['password']:
                user_found = True
                st.session_state.username = username
                st.session_state.role = user['role']
                st.session_state.otp = generate_otp()
                send_otp_email("sricharan320@gmail.com", "geti diky hgdc fiwd", username, st.session_state.otp)
                st.session_state.stage = "otp_verification"
                
                # Initialize user-specific chat history if it doesn't exist
                if f'chat_history_{username}' not in st.session_state:
                    st.session_state[f'chat_history_{username}'] = []
                st.rerun()
                break
        if not user_found:
            st.error("Invalid username or password")

def otp_verification():
    st.title("🔐 OTP Verification")
    st.markdown("### Please enter the OTP sent to your email")
    user_input = st.text_input("Enter OTP:")
    
    if st.button("Verify"):
        if verify_otp(user_input, st.session_state.otp):
            st.session_state.logged_in = True
            st.session_state.stage = "dashboard"
            st.success("✅ OTP Verified Successfully")
            st.rerun()
        else:
            st.error("❌ Invalid OTP. Please try again.")

def extract_text_from_file(uploaded_file):
    file_text = ""
    if uploaded_file.type == "application/pdf":
        pdf_reader = PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            file_text += page.extract_text()
    elif uploaded_file.type == "text/plain":
        file_text = str(uploaded_file.read(), "utf-8")
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(uploaded_file)
        for paragraph in doc.paragraphs:
            file_text += paragraph.text
    return file_text

def generate_audio(text, lang):
    if not text:
        raise ValueError("No text to speak.")
    languages = {"English": "en", "Hindi": "hi", "Marathi": "mr", "Bengali": "bn", "Gujarati": "gu",
                 "Telugu": "te", "Tamil": "ta", "Malayalam": "ml", "Kannada": "kn", "Punjabi": "pa"}
    lang_code = languages.get(lang, "en")
    tts = gTTS(text=text, lang=lang_code)
    audio_io = io.BytesIO()
    tts.write_to_fp(audio_io)
    audio_io.seek(0)
    return audio_io

def translate_text(text, src, target_lang):
    chunks = textwrap.wrap(text, 500)
    translated_chunks = [GoogleTranslator(source=src, target=target_lang).translate(chunk) for chunk in chunks]
    return ' '.join(translated_chunks)

# def dashboard():
#     st.sidebar.markdown("<h1 style='color:white;'>📄 Available Documents</h1>", unsafe_allow_html=True)
#     st.sidebar.write(f"👋 Welcome, {st.session_state.username}!")
    
#     data = load_login_data()
#     role = st.session_state.role
#     accessible_docs = data['roles'].get(role, [])
#     documents = data['documents']
    
#     if not accessible_docs:
#         st.sidebar.warning("No documents available for your role.")
#         return
    
#     for doc in accessible_docs:
#         doc_name = doc
#         doc_path = documents.get(doc, None)
#         if doc_path and os.path.exists(doc_path):
#             with open(doc_path, "rb") as file:
#                 st.sidebar.download_button(
#                     label=f"📥 Download {doc_name}",
#                     data=file,
#                     file_name=doc_name,
#                     mime="application/octet-stream"
#                 )
#         else:
#             st.sidebar.error(f"❌ Document {doc_name} not found.")
    
#     chatbot_interface()

#     if st.sidebar.button("🚪 Sign Out"):
#         st.session_state.logged_in = False
#         st.session_state.stage = "login"
#         st.rerun()

def chatbot_interface():
    st.title("💬 Document-based Q&A Chatbot")
    prev_lang='en'
    username = st.session_state.username

    if f'chat_history_{username}' not in st.session_state:
        st.session_state[f'chat_history_{username}'] = []
    
    # Uploading documents
    uploaded_files = st.file_uploader("📤 Upload your documents", accept_multiple_files=True, type=["txt", "pdf", "docx"])

    documents = []
    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_content = extract_text_from_file(uploaded_file)
            doc_obj = Document(text=file_content)
            documents.append(doc_obj)
        st.success(f"✅ Loaded {len(documents)} document(s).")

    # Set up Together LLM API and embeddings
    api_key = 'dd0fd0c2e4d007fb0deb98a42b41824008a87eb02b716e1eee6a32551253ba9d'
    llm = TogetherLLM(model="meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo", api_key=api_key)
    embeddings = HuggingFaceEmbeddings(api_key='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', model_name="sentence-transformers/all-mpnet-base-v2")

    if st.button("🔄 Process Documents") and documents:
        with st.spinner('Processing...'):
            # Split documents into chunks
            text_splitter = SentenceSplitter(chunk_size=1024, chunk_overlap=200)
            nodes = text_splitter.get_nodes_from_documents(documents, show_progress=True)

            # Define transformations
            transformations = [SentenceSplitter(chunk_size=1024)]
            Settings.transformations = transformations

            # Create VectorStoreIndex and persist storage context
            vector_index = VectorStoreIndex.from_documents(documents, embed_model=embeddings, llm=llm, transformations=transformations)
            vector_index.storage_context.persist(persist_dir="./storage_mini")
        
        st.success("✅ Documents processed")


    st.subheader("📜 Chat History")
    for message in st.session_state[f'chat_history_{username}']:
        st.write(f"{message['role']}: {message['text']}")
        if 'translated' in message:
            st.write(f"Translated ({message['translated']['language']}): {message['translated']['text']}")

    query = st.text_input("💡 Ask a question to the chatbot:")
    if query:
        st.session_state[f'chat_history_{username}'].append({"role": "User", "text": query})

        # Process the user query
        if documents:
            # Load the index and create query engine
            storage_context = StorageContext.from_defaults(persist_dir="./storage_mini")
            index = VectorStoreIndex.from_documents(documents, embed_model=embeddings)  # Adjust to load the index
            query_engine = index.as_query_engine(llm=llm)
            response = query_engine.query(query)
            response_text = response.response

            st.session_state[f'chat_history_{username}'].append({"role": "Bot", "text": response_text})
            
            st.subheader("🌍 Translate response")

            selected_language = st.radio("Choose a language to translate:", list(LANGUAGE_CODES.values()))
            if selected_language:
                lang_code = list(LANGUAGE_CODES.keys())[list(LANGUAGE_CODES.values()).index(selected_language)]
                translated_text = translate_text(response_text, prev_lang, lang_code)
                st.write(f"Translated ({selected_language}): {translated_text}")
                prev_lang=lang_code


                translated_audio_io = generate_audio(translated_text, lang_code)
                st.audio(translated_audio_io, format='audio/mp3')

                st.session_state[f'chat_history_{username}'][-1]['translated'] = {
                'language': selected_language,
                'text': translated_text
            }

# def main():
#     if 'stage' not in st.session_state:
#         st.session_state.stage = "login"

#     if st.session_state.stage == "login":
#         login()
#     elif st.session_state.stage == "otp_verification":
#         otp_verification()
#     elif st.session_state.logged_in:
#         dashboard()

# if __name__ == "__main__":
#     main()
