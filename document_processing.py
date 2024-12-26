import streamlit as st
import os
import random
import smtplib
import string
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import io
from PyPDF2 import PdfReader
import docx
from gtts import gTTS
from langdetect import detect
from deep_translator import GoogleTranslator
from llama_index.core import VectorStoreIndex, StorageContext, Document
from llama_index.llms.together import TogetherLLM
from llama_index.core.node_parser import SentenceSplitter
from langchain_community.embeddings import HuggingFaceInferenceAPIEmbeddings
import textwrap
from llama_index.core import Settings
from utils import load_login_data
import translation as t

LANGUAGE_CODES = {
    'en': 'English', 'hi': 'Hindi', 'mr': 'Marathi', 'bn': 'Bengali',
    'gu': 'Gujarati', 'te': 'Telugu', 'ta': 'Tamil', 'ml': 'Malayalam',
    'kn': 'Kannada', 'pa': 'Punjabi', 'or': 'Odia'
}

def generate_otp(length=6):
    digits = string.digits
    otp = ''.join(random.choice(digits) for _ in range(length))
    return otp

def send_otp_email(sender_email, app_password, receiver_email, otp):
    subject = "Your OTP Code"
    body = f"Your OTP code is {otp}. Please use this code to complete your verification."
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, app_password)
        server.sendmail(sender_email, receiver_email, msg.as_string())
        server.quit()
        st.success(f"OTP sent successfully to {receiver_email}")
    except Exception as e:
        st.error(f"Failed to send OTP. Error: {str(e)}")

def extract_text_from_file(uploaded_file):
    file_text = ""
    if uploaded_file.type == "application/pdf":
        pdf_reader = PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            file_text += page.extract_text()
    elif uploaded_file.type == "text/plain":
        file_text = str(uploaded_file.read(), "utf-8")
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(uploaded_file)
        for paragraph in doc.paragraphs:
            file_text += paragraph.text
    return file_text

def dashboard():
    st.sidebar.markdown("<h1 style='color:white;'>📄 Available Documents</h1>", unsafe_allow_html=True)
    st.sidebar.write(f"👋 Welcome, {st.session_state.username}!")
    
    data = load_login_data()
    role = st.session_state.role
    accessible_docs = data['roles'].get(role, [])
    documents = data['documents']
    
    if not accessible_docs:
        st.sidebar.warning("No documents available for your role.")
        return
    
    for doc in accessible_docs:
        doc_name = doc
        doc_path = documents.get(doc, None)
        if doc_path and os.path.exists(doc_path):
            with open(doc_path, "rb") as file:
                st.sidebar.download_button(
                    label=f"📥 Download {doc_name}",
                    data=file,
                    file_name=doc_name,
                    mime="application/octet-stream"
                )
        else:
            st.sidebar.error(f"❌ Document {doc_name} not found.")
    
    chatbot_interface()

    if st.sidebar.button("🚪 Sign Out"):
        st.session_state.logged_in = False
        st.session_state.stage = "login"
        st.rerun()

def chatbot_interface():
    st.title("💬 Document-based Q&A Chatbot")
    prev_lang='en'
    username = st.session_state.username

    if f'chat_history_{username}' not in st.session_state:
        st.session_state[f'chat_history_{username}'] = []
    
    # Uploading documents
    uploaded_files = st.file_uploader("📤 Upload your documents", accept_multiple_files=True, type=["txt", "pdf", "docx"])

    documents = []
    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_content = extract_text_from_file(uploaded_file)
            doc_obj = Document(text=file_content)
            documents.append(doc_obj)
        st.success(f"✅ Loaded {len(documents)} document(s).")

    # Set up Together LLM API and embeddings
    api_key = 'dd0fd0c2e4d007fb0deb98a42b41824008a87eb02b716e1eee6a32551253ba9d'
    llm = TogetherLLM(model="meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo", api_key=api_key)
    embeddings = HuggingFaceInferenceAPIEmbeddings(api_key='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', model_name="sentence-transformers/all-mpnet-base-v2")

    if st.button("🔄 Process Documents") and documents:
        with st.spinner('Processing...'):
            # Split documents into chunks
            text_splitter = SentenceSplitter(chunk_size=1024, chunk_overlap=200)
            nodes = text_splitter.get_nodes_from_documents(documents, show_progress=True)

            # Define transformations
            transformations = [SentenceSplitter(chunk_size=1024)]
            Settings.transformations = transformations

            # Create VectorStoreIndex and persist storage context
            vector_index = VectorStoreIndex.from_documents(documents, embed_model=embeddings, llm=llm, transformations=transformations)
            vector_index.storage_context.persist(persist_dir="./storage_mini")
        
        st.success("✅ Documents processed")


    st.subheader("📜 Chat History")
    for message in st.session_state[f'chat_history_{username}']:
        st.write(f"{message['role']}: {message['text']}")
        if 'translated' in message:
            st.write(f"Translated ({message['translated']['language']}): {message['translated']['text']}")

    query = st.text_input("💡 Ask a question to Dynamo AI:")
    if query:
        st.session_state[f'chat_history_{username}'].append({"role": "User", "text": query})

        # Process the user query
        if documents:
            # Load the index and create query engine
            storage_context = StorageContext.from_defaults(persist_dir="./storage_mini")
            index = VectorStoreIndex.from_documents(documents, embed_model=embeddings)  # Adjust to load the index
            query_engine = index.as_query_engine(llm=llm)
            response = query_engine.query(query)
            response_text = response.response

            st.session_state[f'chat_history_{username}'].append({"role": "Bot", "text": response_text})
            
            st.subheader("🌍 Translate response")

            selected_language = st.radio("Choose a language to translate:", list(LANGUAGE_CODES.values()))
            if selected_language:
                lang_code = list(LANGUAGE_CODES.keys())[list(LANGUAGE_CODES.values()).index(selected_language)]
                translated_text = t.translate_text(response_text, prev_lang, lang_code)
                st.write(f"Translated ({selected_language}): {translated_text}")
                prev_lang=lang_code


                translated_audio_io = t.generate_audio(translated_text, lang_code)
                st.audio(translated_audio_io, format='audio/mp3')

                st.session_state[f'chat_history_{username}'][-1]['translated'] = {
                'language': selected_language,
                'text': translated_text
            }